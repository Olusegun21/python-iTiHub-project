# Writing Paragraph
# Install the pip3 with the docx files through terminal
# Heading and About section
# Work Experiences
# Skills and Footer
# Text to speech
# Requirement.txt
#GitHub

from docx import Document
from docx.shared import Inches
import pyttsx3

def  speak(text):
     pyttsx3.speak(text)



document = Document()

# Adding Picture to the document
document.add_picture(
     "olu pix.jpg",
     width=Inches(2.0))
# name, phone number and email details
name = input("What is your name? ")
speak("Hello " + name + "how are you today?")
phone_number = input("What is your phone number? ")
email = input("What is your email address? ")

# Adding paragraph to the document
document.add_paragraph(
     name + " | " + phone_number + " | " + email)

# about me
document.add_heading("About me")
about_me = input("Tell me about yourself? ")
document.add_paragraph(about_me)

# Work Experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter company's name ")
from_date = input("From Date ")
to_date = input("To Date ")

# This is how you add text to existing paragraph
p.add_run(company + " ").bold = True
p.add_run(from_date + " - " + to_date + "\n").italic = True

experience_details = input(
     "Describe your experience at " + company)
p.add_run(experience_details)

# More experiences
while True:
     has_more_experiences = input(
          "Do you have more experiences? Yes or No ")
     if has_more_experiences.lower() == 'yes':
          p = document.add_paragraph()

          company = input("Enter company's name ")
          from_date = input("From Date ")
          to_date = input("To Date ")

          # This is how you add text to existing paragraph

          p.add_run(company + " ").bold = True
          p.add_run(from_date + "-" + to_date + "\n").italic = True

          experience_details = input(
               "Describe your experience at " + company + " ")
          p.add_run(experience_details)
     else:
          break

# Skills
document.add_heading("Skills")
skills = input("What are your skills? ")
h = document.add_paragraph(skills)
h.style = "List Bullet"

while True:
     has_more_skills = input(
          "Do you have any other skills you want to add? Yes or No ")
     if has_more_skills.lower() == 'yes':
          skills = input("Add more skills? ")
          h = document.add_paragraph(skills)
          h.style = "List Bullet"

     else:
          break

# Footer
section = document.section[0]
footer = section.footer
g = footer.paragraph[0]
g.text = "CV generated using amigoscode video tutorials"

document.save("cv.docx")



